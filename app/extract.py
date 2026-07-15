import io
import os
import time
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure Gemini API
api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from paragraphs and tables inside a DOCX file using python-docx.
    """
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    
    # 1. Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)
            
    # 2. Extract text from tables, joining cells with " | "
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text for cell in row.cells]
            row_text = " | ".join(row_cells)
            if row_text.strip():
                text_parts.append(row_text)
                
    return "\n".join(text_parts)

def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Decodes text bytes using UTF-8, with a fallback to ignore decode errors if any.
    """
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("utf-8", errors="ignore")

def extract_text_from_image(file_bytes: bytes, mime_type: str) -> str:
    """
    Extracts text from JPG/PNG images using Gemini Vision.
    Implements a 3-attempt linear backoff retry on API failures.
    """
    image_data = {
        "mime_type": mime_type,
        "data": file_bytes
    }
    
    prompt = "Extract all text from this image exactly as written. If it's handwritten notes, transcribe them accurately. Return only the extracted text, no commentary."
    
    max_attempts = 3
    delay = 1.0  # seconds between retries
    
    for attempt in range(max_attempts):
        try:
            model = genai.GenerativeModel("models/gemini-flash-lite-latest")
            response = model.generate_content([image_data, prompt])
            return response.text
        except Exception as e:
            print(f"Gemini image text extraction attempt {attempt + 1} failed: {e}")
            if attempt == max_attempts - 1:
                raise e
            time.sleep(delay * (attempt + 1))
            
    raise RuntimeError("Failed to extract text from image after retries")
